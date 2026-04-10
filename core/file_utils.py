"""
File & Resume Utilities
=======================

Provides helper functions for ingesting and parsing standard resume file formats
(PDF, DOCX, DOC, TXT) into raw text strings that can be securely handled by 
the Semantic Matcher and TF-IDF matching engine.
"""

import os
import fitz  # PyMuPDF
from docx import Document

def extract_text_from_file(file_path):
    """
    Safely extracts and decodes plain text from a given file path.
    
    Supported Formats:
    - `.pdf`: Uses PyMuPDF (fitz) to extract text layer page by page.
    - `.docx`: Uses `python-docx` to iterate and join paragraph text.
    - `.doc`: Provides a degraded fallback/warning (as pure .doc requires heavy external tools).
    - Other/`.txt`: Raw utf-8 text read.
    
    Args:
        file_path (str): Absolute or relative path to the physical resume file on disk.
        
    Returns:
        str: The fully extracted raw text, or an empty string/warning if extraction fails.
    """
    if not os.path.exists(file_path):
        return ""

    ext = os.path.splitext(file_path)[1].lower()
    
    try:
        if ext == ".pdf":
            text = ""
            with fitz.open(file_path) as doc:
                for page in doc:
                    text += page.get_text()
            return text
        
        elif ext == ".docx":
            doc = Document(file_path)
            return "\n".join([para.text for para in doc.paragraphs])
            
        elif ext == ".doc":
            # For .doc, we usually need external tools on Windows, or simplified text extraction.
            # PyMuPDF can sometimes handle .doc via mutation, but Document(docx) won't work.
            # We'll suggest the user convert to .docx for best results, but try a basic read.
            try:
                # If we have antiword or similar, we could use it, but for now, we'll return a warning.
                return "Warning: .doc files are older and less reliable for extraction. Please convert to .docx if this fails."
            except:
                return ""
        
        else:
            # Try to read as plain text
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
    except Exception as e:
        print(f"Error extracting text from {file_path}: {e}")
        return ""

if __name__ == "__main__":
    # Quick test
    import sys
    if len(sys.argv) > 1:
        print(extract_text_from_file(sys.argv[1]))
